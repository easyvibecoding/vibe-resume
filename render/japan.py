"""Japan rirekisho (履歴書) DOCX renderer — JIS Z 8303 simplified grid.

Why this is its own module (and not a Jinja template like every other locale):
the Japanese rirekisho is a fixed-grid form, not flowing prose. Reviewers
(especially in traditional industries) expect specific cell positions —
photo top-right, ふりがな row above the 氏名 row, 学歴/職歴 in a 3-column
year/month/content grid, etc. Markdown→pandoc cannot reliably produce
that layout, so we drive python-docx tables directly.

Companion: a `職務経歴書` (shokumu keirekisho) — the work-history narrative
document that pairs with the rirekisho — is rendered separately as a
flowing markdown template (`resume.ja_JP_shokumu.md.j2`).

This MVP covers:
- 写真欄 (photo cell, optional, takes `profile.photo_path`)
- 氏名 (kanji) + ふりがな
- 生年月日 + 年齢 + 性別
- 現住所 + 電話 + メール
- 学歴・職歴 grid (sourced from profile.education / profile.experience)
- 免許・資格 (profile.certifications)
- 志望動機・自己PR (profile.motivation_ja_JP or summary_ja_JP)

Not yet implemented (intentional MVP scope):
- 通勤時間 / 扶養家族 / 配偶者 cells (legacy fields, often dropped in
  modern tech-industry rirekisho)
- 本人希望記入欄 free-form box
- Stamp box (印鑑欄) — rare on softcopy submissions
- Strict mm-precision cell sizing (we approximate with Cm units)

To extend, add fields to `profile.yaml` (schema is `extra="allow"`) and
new cells in `_personal_info_table` / `_motivation_section`.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from rich.console import Console

console = Console()

DEFAULT_FONT = "Yu Mincho"  # macOS / Office bundles this; falls back gracefully
LABEL_FONT_SIZE = Pt(9)
BODY_FONT_SIZE = Pt(10.5)


# -- low-level cell helpers --------------------------------------------------

def _set_cell_text(cell, text: str, *, bold: bool = False, size: Pt = BODY_FONT_SIZE,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Replace a cell's text with a single styled run, leaving no leftover paragraphs."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = size
    run.font.name = DEFAULT_FONT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_label(cell, text: str) -> None:
    _set_cell_text(cell, text, bold=True, size=LABEL_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)


def _normalize_date(value: Any) -> tuple[str, str, str]:
    """Return (year, month, day-or-empty) strings from `2024-02-14` / `2024-02` / `2024`."""
    if not value:
        return "", "", ""
    s = str(value).strip()
    parts = s.replace("/", "-").replace(".", "-").split("-")
    parts = [p for p in parts if p]
    while len(parts) < 3:
        parts.append("")
    return parts[0], parts[1].lstrip("0"), parts[2].lstrip("0")


def _calc_age_jp(dob: str | None) -> str:
    """Naive age calc — Japan uses date-of-birth-based age (満年齢)."""
    if not dob:
        return ""
    try:
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日"):
            try:
                d = datetime.strptime(dob.strip(), fmt)
                break
            except ValueError:
                continue
        else:
            return ""
    except Exception:
        return ""
    today = datetime.today()
    age = today.year - d.year - ((today.month, today.day) < (d.month, d.day))
    return f"({age}歳)"


def _localized_field(profile: dict[str, Any], key: str, default: str = "") -> str:
    """Profile loader is `extra='allow'` so `<key>_ja_JP` lives next to `<key>`."""
    return profile.get(f"{key}_ja_JP") or profile.get(key) or default


# -- top-level table builders -----------------------------------------------

def _personal_info_table(doc, profile: dict[str, Any], photo_path: Path | None) -> None:
    # 6 rows × 4 cols. Column 4 is the photo (merged top-to-bottom).
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    widths = (Cm(2.5), Cm(6.5), Cm(4.0), Cm(3.5))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    # Photo cell — merge column 4 across all 6 rows
    photo_cell = table.cell(0, 3)
    for r in range(1, 6):
        photo_cell.merge(table.cell(r, 3))
    if photo_path and Path(photo_path).exists():
        photo_cell.text = ""
        run = photo_cell.paragraphs[0].add_run()
        try:
            run.add_picture(str(photo_path), width=Cm(3.0), height=Cm(4.0))
        except Exception as e:
            console.print(f"[yellow]photo embed failed:[/yellow] {e}")
            _set_cell_text(photo_cell, "写真貼付欄\n40mm × 30mm",
                           size=LABEL_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _set_cell_text(photo_cell, "写真貼付欄\n40mm × 30mm",
                       size=LABEL_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 0: ふりがな label + value (cols 0-2)
    _set_label(table.cell(0, 0), "ふりがな")
    furi_cell = table.cell(0, 1).merge(table.cell(0, 2))
    _set_cell_text(furi_cell, profile.get("furigana_ja_JP", ""), size=LABEL_FONT_SIZE)

    # Row 1: 氏名
    _set_label(table.cell(1, 0), "氏 名")
    name_cell = table.cell(1, 1).merge(table.cell(1, 2))
    _set_cell_text(name_cell, _localized_field(profile, "name"), bold=True, size=Pt(14))

    # Row 2: 生年月日 + 性別 (single row split across 3 cols)
    _set_label(table.cell(2, 0), "生年月日")
    dob = profile.get("dob") or profile.get("dob_ja_JP", "")
    y, m, d = _normalize_date(dob)
    dob_str = f"{y}年{m}月{d}日 {_calc_age_jp(dob)}" if y else "____年__月__日"
    _set_cell_text(table.cell(2, 1), dob_str)
    _set_label(table.cell(2, 2), f"性別 {profile.get('gender', '')}")

    # Row 3: 現住所 (ふりがな)
    _set_label(table.cell(3, 0), "ふりがな")
    addr_furi = table.cell(3, 1).merge(table.cell(3, 2))
    _set_cell_text(addr_furi, profile.get("location_furigana_ja_JP", ""), size=LABEL_FONT_SIZE)

    # Row 4: 現住所
    _set_label(table.cell(4, 0), "現住所")
    addr_cell = table.cell(4, 1).merge(table.cell(4, 2))
    _set_cell_text(addr_cell, _localized_field(profile, "location"))

    # Row 5: 電話 + Email split
    _set_label(table.cell(5, 0), "連絡先")
    contact_cell = table.cell(5, 1).merge(table.cell(5, 2))
    parts = []
    if profile.get("phone"):
        parts.append(f"TEL: {profile['phone']}")
    if profile.get("email"):
        parts.append(f"Email: {profile['email']}")
    _set_cell_text(contact_cell, "　".join(parts))


def _history_table(doc, profile: dict[str, Any]) -> None:
    education = [e for e in (profile.get("education") or []) if e.get("school")]
    experience = [e for e in (profile.get("experience") or []) if e.get("company")]

    rows_needed = 1 + 1 + len(education) + 1 + len(experience) + 1  # header + section labels + entries + 以上
    table = doc.add_table(rows=rows_needed, cols=3)
    table.style = "Table Grid"

    widths = (Cm(2.0), Cm(2.0), Cm(12.5))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    # Header row
    _set_label(table.cell(0, 0), "年")
    _set_label(table.cell(0, 1), "月")
    _set_label(table.cell(0, 2), "学歴・職歴")

    idx = 1
    # 学歴 section header (centered, content column)
    _set_cell_text(table.cell(idx, 2), "学歴", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    idx += 1
    for ed in education:
        y, m, _ = _normalize_date(ed.get("year"))
        _set_cell_text(table.cell(idx, 0), y)
        _set_cell_text(table.cell(idx, 1), m)
        school = ed.get("school_ja_JP") or ed.get("school", "")
        degree = ed.get("degree_ja_JP") or ed.get("degree", "")
        _set_cell_text(table.cell(idx, 2), f"{school}　{degree}　卒業")
        idx += 1

    # 職歴 section header
    _set_cell_text(table.cell(idx, 2), "職歴", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    idx += 1
    for e in experience:
        y, m, _ = _normalize_date(e.get("start"))
        _set_cell_text(table.cell(idx, 0), y)
        _set_cell_text(table.cell(idx, 1), m)
        company = e.get("company_ja_JP") or e.get("company", "")
        title = e.get("title_ja_JP") or e.get("title", "")
        _set_cell_text(table.cell(idx, 2), f"{company}　入社（{title}）")
        idx += 1

    # 以上
    _set_cell_text(table.cell(idx, 2), "以上", align=WD_ALIGN_PARAGRAPH.RIGHT)


def _certifications_table(doc, profile: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("【免許・資格】")
    run.bold = True
    run.font.name = DEFAULT_FONT

    certs = profile.get("certifications") or []
    rows = max(1, len(certs))
    table = doc.add_table(rows=rows, cols=3)
    table.style = "Table Grid"
    widths = (Cm(2.0), Cm(2.0), Cm(12.5))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    if not certs:
        _set_cell_text(table.cell(0, 2), "（特になし）", align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    for i, c in enumerate(certs):
        y, m, _ = _normalize_date(c.get("year"))
        _set_cell_text(table.cell(i, 0), y)
        _set_cell_text(table.cell(i, 1), m)
        _set_cell_text(table.cell(i, 2), c.get("name", ""))


def _motivation_section(doc, profile: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("【志望動機・自己PR】")
    run.bold = True
    run.font.name = DEFAULT_FONT

    body = (
        profile.get("motivation_ja_JP")
        or profile.get("summary_ja_JP")
        or profile.get("summary")
        or "（記入してください）"
    )
    body_p = doc.add_paragraph(body)
    for r in body_p.runs:
        r.font.name = DEFAULT_FONT
        r.font.size = BODY_FONT_SIZE


# -- public entry -----------------------------------------------------------

def render_rirekisho(profile: dict[str, Any], out_path: Path,
                     photo_path: Path | None = None) -> Path:
    """Render a JIS-style 履歴書 DOCX to `out_path`. Returns the path."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = BODY_FONT_SIZE

    # Title (centered)
    h = doc.add_heading("履歴書", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date line (top-right)
    today = datetime.today()
    date_str = profile.get("rirekisho_date_ja_JP") or f"{today.year}年{today.month}月{today.day}日 現在"
    p = doc.add_paragraph(date_str)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = DEFAULT_FONT

    _personal_info_table(doc, profile, photo_path)
    doc.add_paragraph()
    _history_table(doc, profile)
    doc.add_paragraph()
    _certifications_table(doc, profile)
    doc.add_paragraph()
    _motivation_section(doc, profile)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path
