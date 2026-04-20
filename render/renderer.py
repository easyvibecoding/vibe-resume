"""Render resume to Markdown / DOCX / PDF and snapshot via git."""
from __future__ import annotations

import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape
from rich.console import Console

from collections import Counter

from core.aggregator import load_groups, load_observed_summary, load_window_stats
from core.schema import load_profile
from core.tech_canonical import category_label, group_by_category, source_display
from core.versioning import snapshot
from render.i18n import format_date, format_date_range, get_locale, localized, resolve_locale

console = Console()
ROOT = Path(__file__).parent.parent


def _history_path(cfg: dict[str, Any]) -> Path:
    p = Path(cfg.get("render", {}).get("output_dir") or "data/resume_history")
    if not p.is_absolute():
        p = ROOT / p
    p.mkdir(parents=True, exist_ok=True)
    return p


def _next_version(dir_: Path) -> int:
    import re

    pat = re.compile(r"resume_v(\d+)")
    existing: list[int] = []
    for p in dir_.glob("resume_v*.md"):
        m = pat.match(p.stem)
        if m:
            existing.append(int(m.group(1)))
    return (max(existing) + 1) if existing else 1


def _pick_template(env: Environment, locale_key: str) -> str:
    """Return template filename, preferring `resume.<locale>.md.j2` over the default."""
    candidate = f"resume.{locale_key}.md.j2"
    try:
        env.get_template(candidate)
        return candidate
    except Exception:
        return "resume.md.j2"


def _render_md(cfg: dict[str, Any], tailor: str | None, locale: str | None = None) -> tuple[str, dict]:
    tpl_dir = Path(cfg.get("render", {}).get("templates_dir") or "render/templates")
    if not tpl_dir.is_absolute():
        tpl_dir = ROOT / tpl_dir
    env = Environment(
        loader=FileSystemLoader(str(tpl_dir)),
        autoescape=select_autoescape(disabled_extensions=("md", "j2")),
        trim_blocks=True,
        lstrip_blocks=True,
        keep_trailing_newline=True,
    )
    locale_key = resolve_locale(locale or cfg.get("render", {}).get("locale"))
    locale_meta = get_locale(locale_key)

    # register locale-aware filters so templates can call {{ d | date }}
    # and {{ entry | localized('bullets') }} without importing anything.
    env.filters["date"] = lambda v: format_date(v, locale_key)
    env.filters["date_range"] = lambda start, end: format_date_range(start, end, locale_key)
    env.filters["localized"] = lambda obj, key: localized(obj, key, locale_key)

    profile = load_profile(ROOT / "profile.yaml")
    groups = load_groups()
    raw_groups = [g.model_dump(mode="json") for g in groups]

    skills: set[str] = set()
    for g in groups:
        skills.update(g.tech_stack)
    skills_list = sorted(skills)
    skills_grouped = group_by_category(skills_list)

    if not raw_groups:
        timespan_start = timespan_end = datetime.now().strftime("%Y-%m-%d")
    else:
        timespan_start = min(g["first_activity"][:10] for g in raw_groups)
        timespan_end = max(g["last_activity"][:10] for g in raw_groups)

    # AI tool usage overview
    src_counter: Counter[str] = Counter()
    for g in groups:
        for s in g.sources:
            src_counter[s.value] += 1
    total_groups = len(groups) or 1
    ai_overview = [
        {
            "tool": source_display(src),
            "projects": cnt,
            "percent": cnt * 100 // total_groups,
        }
        for src, cnt in src_counter.most_common()
    ]
    total_sessions = sum(g.total_sessions for g in groups)

    # top capabilities across all projects (category-level)
    cap_counter: Counter[str] = Counter()
    for g in groups:
        for cat, n in (g.category_counts or {}).items():
            cap_counter[cat] += n
    top_capabilities = [
        category_label(c)
        for c, _ in cap_counter.most_common(8)
        if c != "fullstack"
    ][:6]

    # humanize per-group headlines (turn internal slugs into display labels)
    headline_map = {
        "frontend": "Frontend",
        "backend": "Backend",
        "fullstack": "Full-stack",
        "database": "Database",
        "devops": "DevOps",
        "deployment": "Deployment",
        "bug-fix": "Bug fixes",
        "feature": "Features",
        "refactor": "Refactoring",
        "testing": "Testing",
        "ui-design": "UI",
        "docs": "Docs",
        "performance": "Performance",
        "security": "Security",
        "data-ml": "Data/ML",
        "api-integration": "API integration",
        "agent-tooling": "Agent tooling",
        "research": "Research",
    }
    for g in raw_groups:
        hl = g.get("headline") or ""
        for k, v in headline_map.items():
            hl = hl.replace(f"{k} ", f"{v} ")
        g["headline"] = hl

    observed = load_observed_summary() or {}
    window_stats = load_window_stats() or {}

    ctx = {
        "profile": profile.model_dump(),
        "groups": raw_groups,
        "skills": skills_list,
        "skills_grouped": skills_grouped,
        "ai_overview": ai_overview,
        "total_sessions": total_sessions,
        "top_capabilities": top_capabilities,
        "observed_summary": observed.get("summary"),
        "window_stats": window_stats,
        "tailored_for": tailor,
        "timespan_start": timespan_start,
        "timespan_end": timespan_end,
        "locale": locale_meta,
        "T": locale_meta["headings"],
    }
    tpl_name = _pick_template(env, locale_meta["_key"])
    tpl = env.get_template(tpl_name)
    ctx["_tpl_name"] = tpl_name
    return tpl.render(**ctx), ctx


def _hide_table_borders(table) -> None:
    """Strip all visible borders from a python-docx table — used for layout tables."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _render_docx_header(doc, profile: dict, photo_path: str | None, photo_expected: bool, locale_key: str) -> None:
    """Top-of-page header. Uses an invisible 1x2 table when a photo is expected
    and present (so the photo sits in the true top-right), otherwise a flowing
    name + title + contacts block."""
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    contacts = [profile.get(k) for k in ("email", "phone", "location", "linkedin", "github", "website")]
    contacts = [c for c in contacts if c]
    has_photo = bool(photo_path) and Path(photo_path).exists()

    if photo_expected and has_photo:
        # 1×2 invisible table: left = name+title+contacts, right = photo
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _hide_table_borders(table)

        left = table.cell(0, 0)
        right = table.cell(0, 1)
        left.width = Cm(13.0)
        right.width = Cm(4.5)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # left cell: name (large), title (italic), contacts
        name_p = left.paragraphs[0]
        name_run = name_p.add_run(profile["name"])
        name_run.bold = True
        name_run.font.size = Pt(20)
        if profile.get("title"):
            tp = left.add_paragraph(profile["title"])
            tp.runs[0].italic = True
        if profile.get("target_role"):
            tr = left.add_paragraph()
            tr.add_run("Target role: ").bold = True
            tr.add_run(profile["target_role"])
        if contacts:
            cp = left.add_paragraph(" · ".join(contacts))
            cp.runs[0].font.size = Pt(9)

        # right cell: photo
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            right_p.add_run().add_picture(str(photo_path), width=Cm(3.0), height=Cm(4.0))
        except Exception as e:
            console.print(f"[yellow]photo embed failed for {locale_key}:[/yellow] {e}")
        return

    # Plain flowing layout (en_US, en_GB, zh_TW etc., or photo-expected locale w/o photo)
    doc.add_heading(profile["name"], level=0)
    if photo_expected and not has_photo:
        console.print(f"[yellow]locale {locale_key} expects a photo but profile.photo_path is empty[/yellow]")
    if profile.get("title"):
        doc.add_paragraph(profile["title"]).italic = True
    if contacts:
        doc.add_paragraph(" · ".join(contacts))
    if profile.get("target_role"):
        p = doc.add_paragraph()
        p.add_run("Target role: ").bold = True
        p.add_run(profile["target_role"])


def _render_docx(md_text: str, ctx: dict, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang TC"
    style.font.size = Pt(11)

    profile = ctx["profile"]
    locale_key = ctx.get("locale", {}).get("_key", "en_US")
    photo_expected = ctx.get("locale", {}).get("photo") == "expected"
    photo_path = profile.get("photo_path")

    _render_docx_header(doc, profile, photo_path, photo_expected, locale_key)

    if profile.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(profile["summary"])

    if ctx["skills"]:
        doc.add_heading("Technical skills", level=1)
        doc.add_paragraph(", ".join(ctx["skills"]))

    doc.add_heading("AI-assisted project work", level=1)
    doc.add_paragraph(
        f"Auto-aggregated from {len(ctx['groups'])} project groups; "
        f"window {ctx['timespan_start']} – {ctx['timespan_end']}"
    ).italic = True
    for i, g in enumerate(ctx["groups"], 1):
        doc.add_heading(f"{i}. {g['name']}", level=2)
        p = doc.add_paragraph()
        p.add_run("Period: ").bold = True
        p.add_run(
            f"{g['first_activity'][:10]} – {g['last_activity'][:10]} ({g['total_sessions']} sessions)"
        )
        p = doc.add_paragraph()
        p.add_run("AI tools: ").bold = True
        p.add_run(", ".join(g["sources"]))
        if g.get("tech_stack"):
            p = doc.add_paragraph()
            p.add_run("Stack: ").bold = True
            p.add_run(", ".join(g["tech_stack"]))
        if g.get("summary"):
            doc.add_paragraph(g["summary"])
        for a in g.get("achievements") or []:
            doc.add_paragraph(a, style="List Bullet")

    section_titles = {
        "experience": "Experience",
        "education": "Education",
        "certifications": "Certifications",
    }
    for section, title in section_titles.items():
        items = profile.get(section) or []
        if not items:
            continue
        doc.add_heading(title, level=1)
        for item in items:
            doc.add_paragraph(str(item))

    doc.save(str(out_path))


def _render_pdf(md_path: Path, out_path: Path) -> bool:
    if not shutil.which("pandoc"):
        console.print("[yellow]pandoc not found; skipping PDF[/yellow]")
        return False
    cmd = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_path),
        "--pdf-engine=xelatex",
        "-V",
        "CJKmainfont=PingFang TC",
        "-V",
        "geometry:margin=1in",
    ]
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        # fall back to weasyprint via HTML
        console.print(f"[yellow]pandoc xelatex failed:[/yellow] {r.stderr[:200]}")
        cmd = ["pandoc", str(md_path), "-o", str(out_path)]
        r = subprocess.run(cmd, capture_output=True, text=True)
        if r.returncode != 0:
            console.print(f"[red]pandoc failed:[/red] {r.stderr[:200]}")
            return False
    return True


def render_draft(
    cfg: dict[str, Any],
    fmt: str = "md",
    tailor: str | None = None,
    locale: str | None = None,
) -> None:
    hist = _history_path(cfg)
    version = _next_version(hist)

    md_text, ctx = _render_md(cfg, tailor, locale=locale)
    locale_key = ctx["locale"]["_key"]
    suffix = "" if locale_key == "en_US" else f"_{locale_key}"
    md_path = hist / f"resume_v{version:03d}{suffix}.md"
    md_path.write_text(md_text)
    console.print(f"[green]✓[/green] {md_path.name}  [dim](locale={locale_key}, tpl={ctx['_tpl_name']})[/dim]")

    written = [md_path]
    if fmt in ("docx", "all"):
        docx_path = hist / f"resume_v{version:03d}{suffix}.docx"
        if locale_key == "ja_JP":
            from render.japan import render_rirekisho

            photo = ctx["profile"].get("photo_path")
            photo_p = Path(photo) if photo else None
            render_rirekisho(ctx["profile"], docx_path, photo_path=photo_p)
        else:
            _render_docx(md_text, ctx, docx_path)
        written.append(docx_path)
        console.print(f"[green]✓[/green] {docx_path.name}")
    if fmt in ("pdf", "all"):
        pdf_path = hist / f"resume_v{version:03d}{suffix}.pdf"
        if _render_pdf(md_path, pdf_path):
            written.append(pdf_path)
            console.print(f"[green]✓[/green] {pdf_path.name}")

    msg = f"resume v{version} ({fmt}, {locale_key})"
    if tailor:
        msg += f" [tailored:{Path(tailor).stem}]"
    sha = snapshot(cfg, written, msg)
    if sha:
        console.print(f"[cyan]snapshot[/cyan] {sha}")
